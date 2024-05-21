import pandas as pd
from docx import Document
import time
import streamlit as st
import os

def read_excel_column(file_path, sheet_name, column_letter):
    df = pd.read_excel(file_path, sheet_name=sheet_name, usecols=column_letter, header=None)
    df = df.fillna('')
    data_list = df.iloc[:, 0].tolist()
    return data_list

def create_replacement_dict(data_list):
    return {f'<f{i}>': str(data) for i, data in enumerate(data_list, start=1)}

def replace_placeholders_in_text(text, replacement_dict):
    for placeholder, value in replacement_dict.items():
        if placeholder in text:
            text = text.replace(placeholder, value)
    return text

def replace_placeholders_in_paragraphs(paragraphs, replacement_dict):
    for para in paragraphs:
        para.text = replace_placeholders_in_text(para.text, replacement_dict)

def replace_placeholders_in_tables(tables, replacement_dict):
    for table in tables:
        for row in table.rows:
            for cell in row.cells:
                replace_placeholders_in_paragraphs(cell.paragraphs, replacement_dict)

def replace_placeholders(doc, replacement_dict):
    replace_placeholders_in_paragraphs(doc.paragraphs, replacement_dict)
    replace_placeholders_in_tables(doc.tables, replacement_dict)

def main():
    st.title("Excel to Word Template Replacer")
    
    excel_file = st.file_uploader("Upload Excel File", type=["xlsx"])
    if excel_file is not None:
        sheet_name = st.text_input("Sheet Name", value="Заявка")
        column_letter = st.text_input("Column Letter", value="B")
        
        # Get the base name of the Excel file and create output Word file path
        excel_file_name = excel_file.name
        base_name = os.path.splitext(excel_file_name)[0]
        output_word_path = base_name + ".docx"
        st.write(f"Output Word file will be saved as: {output_word_path}")

    word_template = st.file_uploader("Upload Word Template", type=["docx"])
    
    if st.button("Run"):
        if excel_file and word_template and sheet_name and column_letter:
            start_time = time.time()
            
            data_list = read_excel_column(excel_file, sheet_name, column_letter)
            replacement_dict = create_replacement_dict(data_list)
            
            doc = Document(word_template)
            replace_placeholders(doc, replacement_dict)
            doc.save(output_word_path)
            
            elapsed_time = time.time() - start_time
            st.success(f"Document saved successfully as {output_word_path}")
            st.info(f"Time taken: {elapsed_time:.2f} seconds")
        else:
            st.error("Please upload both Excel and Word files, and provide necessary details.")

if __name__ == '__main__':
    main()
